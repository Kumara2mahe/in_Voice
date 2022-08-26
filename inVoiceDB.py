"""
This file creates all the tables and databases,
used in the in_Voice APP as class,
and also does the CRUD operations of database by using the methods.
"""

# Importing the required modules to working with database
import sqlite3

# Importing os module to work with files and folders
import os

# Importing a function to open and read the data in the url
from urllib.request import urlopen

# Importing json module to convert the data from the url into json format
import json

# Importing the required modules from the Python-Docx module to create and work with '.docx' files
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Creating a class for doing all the CRUD operations in the 'userCredentials' table on the inVoice App's Database
class UserCredential():

    """
    Class for Registering a New Account,
    which is used for working with InVoice App.
    """

    # Assigning a name for the Directory where .db files are stored
    DB_directory = "./.DB"

    def __init__(self, db_name="in_Voice.db"):

        # Assigning 'userCredentials' table name into a variable
        self.tableName = "userCredentials"

        # Creating a '.DB' Directory for storing the database files
        self.createHiddenDIR(self.DB_directory)

        # Connecting to the App's Database, if not exists creating a new database
        self.inVoice_DB = sqlite3.connect(f"{self.DB_directory}/{db_name}")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'userCredentials' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            firstName TEXT,
                            lastName TEXT,
                            email CHAR NOT NULL UNIQUE,
                            mobileNumber INTEGER,
                            userId CHAR NOT NULL PRIMARY KEY,
                            passWord CHAR
                        )""")

    '''CREATE USER'''

    # Function for creating new record in the 'userCredentials' table
    def createUser(self, firstName, lastName, email, mobileNumber, userId, password):

        # Inserting a new record to the table
        self.co.execute(f"INSERT INTO {self.tableName} VALUES (:firstName, :lastName, :email, :mobileNumber, :userId, :password)",
                        {
                            "firstName": firstName,
                            "lastName": lastName,
                            "email": email,
                            "mobileNumber": mobileNumber,
                            "userId": userId,
                            "password": password
                        })

    '''UPDATE USER PASSWORD'''

    # Function for updating only a specific field from a existing record(user detail) in the 'userCredentials' table
    def changePassword(self, userId, password):

        # Updating only the specific field from the existing record(user detail) in the table which matches the userID
        self.co.execute(f"UPDATE {self.tableName} SET password = :password WHERE userId = :userId",
                        {
                            "password": password,

                            "userId": userId
                        })

    '''GET USER'''

    # Function for querying the 'userCredentials' table for specific record
    def getUser(self, userId):

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Querying the 'userCredentials' table, to pick a matched record
        self.co.execute(
            f"SELECT *, oid FROM {self.tableName} WHERE userId = '{userId}'")

        # Assigning the queryed records into a variable to return as response
        records = self.co.fetchone()

        return records

    '''GET ALL USERS'''

    # Function for getting all the records from the table
    def getallUsers(self):

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Querying the 'userCredentials' table, and picking all the available records
        self.co.execute(f"SELECT *, oid FROM {self.tableName}")

        # Assigning the queryed records into a variable to return as response
        records = self.co.fetchall()

        return records

    '''DELETE USER'''

    # Function for deleting a specific record from the 'userCredentials' table
    def delUser(self, userId):

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Deleting a specific record from the 'userCredentials' table which matches the UserId
        self.co.execute(
            f"DELETE FROM {self.tableName} WHERE userId = '{userId}'")

    '''CREATE HIDDEN DIRECTORY'''

    # Function for Creating a Hidden Directory and Deleting the Directory and its files
    def createHiddenDIR(self, dirName, status=""):

        # Confirming Directory not exists
        if (os.path.exists(dirName) is False):

            # Creating a Directory in the name passed as parameter
            os.mkdir(dirName)

            # Making the Directory Hidden
            os.system("attrib +h " + dirName)

        elif (os.path.exists(dirName) and status == "clearCache"):

            # Deleting the all the files in the Directory and its files
            for item in os.listdir(dirName):
                os.remove(f"{dirName}/{item}")

            # Deleting the empty Directory
            os.rmdir(dirName)

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Creating a class for doing all the CRUD operations in the 'userDetails' table on the User's Database
class UserDetail():

    """
    Class for Creating a New Database,
    which is used to separate all user data separate from each other,
    and also from the main InVoice App.
    """

    # Assigning a name for the Directory where binary files are stored
    cache_directory = "./.__appcache__"

    def __init__(self, userId):

        # Assigning the UserId to a variable so it can be accessed by other functions inside this Class
        self.userId = userId

        # Assigning 'userDetails' table name into a variable
        self.tableName = "userDetails"

        # Connecting to the User's Database, if not exists creating a new database for each user
        self.inVoice_DB = sqlite3.connect(
            f"{UserCredential.DB_directory}/{userId}.db")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'userDetails' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            userId CHAR NOT NULL PRIMARY KEY,
                            email CHAR NOT NULL UNIQUE,
                            firstName TEXT NOT NULL,
                            lastName TEXT NOT NULL,
                            profile BLOB NOT NULL
                        )""")

    '''CONVERT IMAGE TO BINARY DATA'''

    # Function for converting Binary file into a BLOB data
    def convertToBinary(self, imageFile="Images/profile.png"):

        # Opening the 'image' file in binary format for reading the binary data
        with open(imageFile, 'rb') as file:

            # Assigning the read binary data to variable and closing the file
            blobData = file.read()
            file.close()

        return blobData

    '''CONVERT BINARY DATA TO IMAGE'''

    # Function for converting BLOB data into Binary file and saving it to a local folder
    def convertToImage(self, blobData, fileName):

        # Generating the path to file from the passed UserID
        pathToImageFile = f"{self.cache_directory}/{fileName}_image.png"

        count = 0
        while (os.path.exists(pathToImageFile)):

            count += 1

            # Generating the new image path if generated path exists
            pathToImageFile = f"{self.cache_directory}/{fileName}_image{count}.png"

        # Opening a new 'image' file in binary format for writing the binary data
        with open(pathToImageFile, 'wb') as file:

            # Writing the binary data queryied from database and close the file
            file.write(blobData)
            file.close()

        return pathToImageFile

    '''CREATE PROFILE'''

    # Function for creating new record in the 'userDetails' table
    def create(self, userId, email, firstName, lastName):

        # Converting binary file to BLOB data
        image = self.convertToBinary()

        # Inserting a new record to the table
        self.co.execute(f"INSERT INTO {self.tableName} VALUES (:userId, :email, :firstName, :lastName, :profile)",
                        {
                            "userId": userId,
                            "email": email,
                            "firstName": firstName,
                            "lastName": lastName,
                            "profile": image
                        })

    '''GET PROFILE'''

    # Function for querying the 'userDetails' table for specific record
    def get(self):

        # Creating a '.__appcache__' Directory for storing the binary files
        UserCredential.createHiddenDIR(self, dirName=self.cache_directory)

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Querying the 'userDetails' table, to pick the binary data from matched record
        self.co.execute(
            f"SELECT * FROM {self.tableName} WHERE userId = '{self.userId}'")

        # Assigning the queryed record into a variable to use further
        record = self.co.fetchone()

        # Converting the queried Blob data into a '.png' file in a specified path
        pathToProfile = self.convertToImage(record[4], record[0])

        # Assigning the queried list as a dictionary and passing it as a response
        userRecords = {
            "userId": str(record[0]),
            "email": str(record[1]),
            "fullName": str(record[2]) + " " + str(record[3]),
            "profile": str(pathToProfile)
        }

        return userRecords

    '''UPDATE PROFILE PICTURE'''

    # Function for updating the User's profile picture in the 'userDetails' table
    def updateProfilePicture(self, pathToNewPicture):

        # Converting binary file to BLOB data
        image = self.convertToBinary(pathToNewPicture)

        # Updating the existing record in the table which matches to UserId
        self.co.execute(f"UPDATE {self.tableName} SET profile = :profile WHERE userId = :userId",
                        {
                            "profile": image,
                            "userId": self.userId
                        })

        # Deleting the temporary image file passed by User
        os.remove(pathToNewPicture)

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Creating a class for doing all the CRUD operations in the 'userPreferences' table on the inVoice User's Database
class UserPreference():

    """
    Class for creating and updating the User's Preference for working with the App,
    which includes the default for Currency, tax percentage, customer message and due date.
    """

    def __init__(self, userId):

        # Assigning 'userPreferences' table name into a variable
        self.tableName = "userPreferences"

        # Connecting to the App's Database, if not exists creating a new database
        self.inVoice_DB = sqlite3.connect(
            f"{UserCredential.DB_directory}/{userId}.db")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'userPreferences' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            currencyCode TEXT NOT NULL UNIQUE,
                            dueDate TEXT NOT NULL,
                            taxInPercentage INTEGER NOT NULL
                        )""")

    '''CREATE USER DEFAULT PREFERENCES'''

    # Function for creating new record in the 'userPreferences' table
    def createDefault(self):

        # Inserting a new record to the table
        self.co.execute(f"INSERT INTO {self.tableName} VALUES (:currencyCode, :dueDate, :taxInPercentage)",
                        {
                            "currencyCode": "INR",
                            "dueDate": "Nil",
                            "taxInPercentage": 0
                        })

    '''UPDATE USER DEFAULT PREFERENCES'''

    # Function for updating the existing record(user preference) in the 'userPreferences' table
    def updateDefault(self, currencyCode, dueDate, taxInPercentage):

        # Updating the existing record in the table which matches the uniqueID
        self.co.execute(f"UPDATE {self.tableName} SET currencyCode = :currencyCode, dueDate = :dueDate, taxInPercentage = :taxInPercentage WHERE oid = :uniqueID",
                        {
                            "currencyCode": currencyCode,
                            "dueDate": dueDate,
                            "taxInPercentage": taxInPercentage,

                            "uniqueID": "1"
                        })

    '''GET USER PREFERENCES'''

    # Function for getting all the records from the 'userPreferences' table
    def all(self):

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Querying the 'currencyRates' table, and picking all the available records
        self.co.execute(f"SELECT * FROM {self.tableName}")

        # Assigning the queryed records into a variable to return as response
        records = self.co.fetchone()

        return records

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Creating a class for doing all the CRUD operations in the 'currencyRates' table on the inVoice User's Database
class CurrencyRate():

    """
    Class for creating and updating the current Currency Exchange Rate to INR,
    and the json data used in this process is collected from (http://www.floatrates.com).
    """

    # Assigning the website URL into a variable which contains the JSON data for INR
    site_URL = "http://www.floatrates.com/daily/inr.json"

    # Defining a List of Currency Codes and Symbols for specific Countries as a python tuple list
    CODES_and_SYMBOLS = [("aud", "$"),
                         ("cad", "$"),
                         ("chf", "chf"),
                         ("eur", "€"),
                         ("gbp", "£"),
                         ("inr", "₹"),
                         ("jpy", "¥"),
                         ("nzd", "$"),
                         ("usd", "$"),
                         ("zar", "R")]

    # Defining the App's default List for Currency Exchange Rate
    default_currencyRates = [("AUD", "Australian Dollar", "$", 0.02, 55.68),
                             ("CAD", "Canadian Dollar", "$", 0.02, 61.78),
                             ("CHF", "Swiss Franc", "chf", 0.01, 83.49),
                             ("EUR", "Euro", "€", 0.01, 80.73),
                             ("GBP", "U.K. Pound Sterling", "£", 0.01, 96.06),
                             ("INR", "Indian Rupee", "₹", 1.0, 1.0),
                             ("JPY", "Japanese Yen", "¥", 1.70, 0.59),
                             ("NZD", "New Zealand Dollar", "$", 0.02, 50.41),
                             ("USD", "U.S. Dollar", "$", 0.01, 79.34),
                             ("ZAR", "South African Rand", "R", 0.21, 4.82)]

    def __init__(self, userId):

        # Assigning 'currencyRates' table name into a variable
        self.tableName = "currencyRates"

        # Connecting to the App's Database, if not exists creating a new database
        self.inVoice_DB = sqlite3.connect(
            f"{UserCredential.DB_directory}/{userId}.db")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'currencyRates' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            currencyCode TEXT NOT NULL PRIMARY KEY,
                            currencyName TEXT NOT NULL,
                            symbol TEXT NOT NULL,
                            INRvalue REAL NOT NULL,
                            exchangeValue REAL NOT NULL
                        )""")

    '''UPDATE CURRENCY EXCHANGE RATE'''

    # Function for creating new record in the 'currencyRates' table
    def update(self, status="update"):

        try:

            # Opening the website URL
            weburl = urlopen(self.site_URL, timeout=1)
            resultCode = weburl.getcode()

            # Creating a List to hold the Currency Exchange Rate for each Country
            currencyRates = []

            # Checking the result code of the website is 200
            if (resultCode == 200):

                # Reading the data on the website URL and assigning it into a variable
                data = weburl.read()

                # Parsing the JSON data available on the website and storing it on a variable
                jsonData = json.loads(data)

                # Iterating through each Country's currency code
                for code in self.CODES_and_SYMBOLS:

                    # Checking the Country's currency code exists in the JSON data
                    if (jsonData.get(code[0]) is None):

                        # Creating a tuple with the values for the each Country's Currency (Code, name, rate, exchangerate)
                        item = ("INR",
                                "Indian Rupee",
                                code[1],
                                "%.2f" % 1,
                                "%.2f" % 1)

                        # Adding the tuple into the Currency Exchange Rate List as a item
                        currencyRates.append(item)

                    else:

                        # Creating a tuple with the values for the each Country's Currency (Code, name, rate, exchangerate)
                        item = (jsonData.get(code[0])["code"],
                                jsonData.get(code[0])["name"],
                                code[1],
                                "%.2f" % jsonData.get(code[0])["rate"],
                                "%.2f" % jsonData.get(code[0])["inverseRate"])

                        # Adding the tuple into the Currency Exchange Rate List as a item
                        currencyRates.append(item)

        except:

            # Querying the 'currencyRates' table, and picking all the available records(Currencies) if has one
            if (self.all() != []):

                currencyRates = self.all()
            else:

                # Assigning the App's default Currency Exchange Rate List
                currencyRates = self.default_currencyRates

        finally:

            # Iterating through each Country's currency code in the Currency Exchange Rate List
            for item in currencyRates:

                # Confirming the status is 'update'
                if (status == "update"):

                    # Updating the existing record in the table which matches to Currency Code
                    self.co.execute(f"UPDATE {self.tableName} SET INRvalue = :INRvalue, exchangeValue = :exchangeValue WHERE currencyCode = :uniqueID",
                                    {
                                        "INRvalue": float(item[3]),
                                        "exchangeValue": float(item[4]),

                                        "uniqueID": item[0]
                                    })

                elif (status == "create"):

                    # Inserting a new record to the table for every item if status is 'create'
                    self.co.execute(f"INSERT INTO {self.tableName} VALUES (:currencyCode, :currencyName, :symbol, :INRvalue, :exchangeValue)",
                                    {
                                        "currencyCode": item[0],
                                        "currencyName": item[1],
                                        "symbol": item[2],
                                        "INRvalue": float(item[3]),
                                        "exchangeValue": float(item[4])
                                    })

    '''GET ALL CURRENCY EXCHANGE RATES'''

    # Function for getting all the records(Currencies) from the 'currencyRates' table
    def all(self):

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Querying the 'currencyRates' table, and picking all the available records(Currencies)
        self.co.execute(f"SELECT * FROM {self.tableName}")

        # Assigning the queryed records(Currencies) into a variable to return as response
        records = self.co.fetchall()

        return records

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Creating a class for doing all the CRUD operations in the 'dueDates' table on the inVoice User's Database
class DueDate():

    """
    Class for creating and updating the due dates,
    which is in the form of days count.
    """

    # Defining the App's default List for Due Dates dropdown
    default_dueDates = ["Nil", "7", "14", "28", "56", "84"]

    def __init__(self, userId):

        # Assigning 'dueDates' table name into a variable
        self.tableName = "dueDates"

        # Connecting to the App's Database, if not exists creating a new database
        self.inVoice_DB = sqlite3.connect(
            f"{UserCredential.DB_directory}/{userId}.db")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'dueDates' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            dayCount TEXT NOT NULL
                        )""")

    '''ADD DUE DATE'''

    # Function for adding new record in the 'dueDates' table
    def add(self, days=None):

        if (days is None):

            for day in self.default_dueDates:

                # Inserting a new record to the table
                self.co.execute(f"INSERT INTO {self.tableName} VALUES (:dayCount)",
                                {
                                    "dayCount": str(day)
                                })

        elif (days is not None):

            # Inserting a new record to the table
            self.co.execute(f"INSERT INTO {self.tableName} VALUES (:dayCount)",
                            {
                                "dayCount": str(days)
                            })

    '''GET ALL DUE DATE'''

    # Function for getting all the records from the 'dueDates' table
    def all(self):

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Querying the 'dueDates' table, and picking all the available records
        self.co.execute(f"SELECT *, oid FROM {self.tableName}")

        # Assigning the queryed records into a variable to return as response
        records = self.co.fetchall()

        return records

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Creating a class for doing all the CRUD operations in the 'inVoiceDetails' table on the inVoice User's Database
class InVoiceDetail():

    """
    Class for Creating a New Table in the User's Database,
    which contains all the Invoices created by the User for their clients.
    Each Invoice holds the client details, date of purchase, due date , total amount , inVoice number which is unique,
    and also has some additional data.
    """

    def __init__(self, userId):

        # Assigning the UserId to a variable so it can be accessed by other functions inside this Class
        self.userId = userId

        # Assigning 'inVoiceDetails' table name into a variable
        self.tableName = "inVoiceDetails"

        # Connecting to the User's Database, if not exists creating a new database for each user
        self.inVoice_DB = sqlite3.connect(
            f"{UserCredential.DB_directory}/{userId}.db")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'inVoiceDetails' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            invoiceNumber CHAR NOT NULL PRIMARY KEY,
                            clientName TEXT NOT NULL,
                            currencyCode TEXT NOT NULL,
                            currencySymbol TEXT NOT NULL,
                            datePurchased DATE NOT NULL,
                            dueDate DATE NOT NULL,
                            productsPurchased BLOB NOT NULL,
                            customerMessage TEXT,
                            taxInPercentage INTEGER,
                            subTotal REAL NOT NULL,
                            calculatedTAX REAL NOT NULL,
                            totalAmount REAL NOT NULL,
                            balanceAmount REAL NOT NULL,
                            paymentStatus TEXT
                        )""")

    '''CREATE INVOICE'''

    # Function for creating new record in the 'inVoiceDetails' table
    def create(self, inVoiceNo, clientName, currencyCode, currencySymbol, purchaseDate, dueDate, productsPurchased, customerMessage, taxInPercentage, subTotal, calculatedTAX, totalAmount, balanceAmount):

        # Converting a list of tuples to json data
        purchasedProducts = json.dumps(productsPurchased)

        # Setting the value of payment status by checking the Balance Amount
        paymentStatus = "Paid" if (int(balanceAmount) == 0) else "Pending"

        # Inserting a new record to the table
        self.co.execute(f"INSERT INTO {self.tableName} VALUES (:invoiceNumber, :clientName, :currencyCode, :currencySymbol, :datePurchased, :dueDate, :productsPurchased, :customerMessage, :taxInPercentage, :subTotal, :calculatedTAX, :totalAmount, :balanceAmount, :paymentStatus)",
                        {
                            "invoiceNumber":  inVoiceNo,
                            "clientName": clientName,
                            "currencyCode": currencyCode,
                            "currencySymbol": currencySymbol,
                            "datePurchased": purchaseDate,
                            "dueDate": dueDate,
                            "productsPurchased": purchasedProducts,
                            "customerMessage": customerMessage,
                            "taxInPercentage": taxInPercentage,
                            "subTotal": subTotal,
                            "calculatedTAX": calculatedTAX,
                            "totalAmount": totalAmount,
                            "balanceAmount": balanceAmount,
                            "paymentStatus": paymentStatus
                        })

    '''GET INVOICE'''

    # Function for getting specific record(inVoice) from the 'inVoiceDetails' table
    def get(self, inVoiceId):

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Querying the 'inVoiceDetails' table, and picking the matched record(inVoice)
        self.co.execute(
            f"SELECT * FROM {self.tableName} WHERE invoiceNumber = '{inVoiceId}'")

        # Assigning the queryed record(inVoice) into a variable to return as response
        records = self.co.fetchone()

        # Creating a empty python list to store the queryed record
        new_records = []
        for item in records:

            if (records.index(item) != 6):

                # Adding data to the empty list
                new_records.append(item)
            else:

                # Converting the BLOB data to a python list and adding it to the empty list
                new_records.append(json.loads(item))

        return new_records

    '''UPDATE INVOICE'''

    # Function for updating specific record(inVoice) in the 'inVoiceDetails' table
    def update(self, inVoiceId, clientName, currencyCode, currencySymbol, purchaseDate, dueDate, productsPurchased, customerMessage, taxInPercentage, subTotal, calculatedTAX, totalAmount, balanceAmount):

        # Converting a list of tuples to json data
        purchasedProducts = json.dumps(productsPurchased)

        # Setting the value of payment status by checking the Balance Amount
        paymentStatus = "Paid" if (int(balanceAmount) == 0) else "Pending"

        # Updating a specific record(inVoice) in the 'inVoiceDetails' table using its invoiceNumber
        self.co.execute(f"UPDATE {self.tableName} SET clientName = :clientName, currencyCode = :currencyCode, currencySymbol = :currencySymbol, datePurchased = :datePurchased, dueDate = :dueDate, productsPurchased = :productsPurchased, customerMessage = :customerMessage, taxInPercentage = :taxInPercentage, subTotal = :subTotal, calculatedTAX = :calculatedTAX, totalAmount = :totalAmount, balanceAmount = :balanceAmount, paymentStatus = :paymentStatus WHERE invoiceNumber = :inVoiceId",
                        {
                            "clientName": clientName,
                            "currencyCode": currencyCode,
                            "currencySymbol": currencySymbol,
                            "datePurchased": purchaseDate,
                            "dueDate": dueDate,
                            "productsPurchased": purchasedProducts,
                            "customerMessage": customerMessage,
                            "taxInPercentage": taxInPercentage,
                            "subTotal": subTotal,
                            "calculatedTAX": calculatedTAX,
                            "totalAmount": totalAmount,
                            "balanceAmount": balanceAmount,
                            "paymentStatus": paymentStatus,

                            "inVoiceId": inVoiceId
                        })

    '''GET ALL INVOICES'''

    # Function for getting all the records(invoices) from the 'inVoiceDetails' table
    def all(self):

        # Querying the 'inVoiceDetails' table, and picking all the available records(invoices)
        self.co.execute(f"SELECT *, oid FROM {self.tableName}")

        # Assigning the queryed records(invoices) into a variable to return as response
        records = self.co.fetchall()

        return records

    '''DELETE INVOICE DETAILS'''

    # Function for deleting a specific record(inVoice) from the 'inVoiceDetails' table
    def delete(self, userId):

        # Deleting a specific record(inVoice) from the 'inVoiceDetails' table which matches the record(inVoice)
        self.co.execute(
            f"DELETE FROM {self.tableName} WHERE invoiceNumber = '{userId}'")

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Creating a class for doing all the CRUD operations in the 'clientDetails' table on the inVoice User's Database
class ClientDetail():

    """
    Class for Creating a New Table in the User's Database,
    which contains all the Clients data created by the User for using it later on creating InVoice.
    Each record has the individual client details.
    """

    def __init__(self, userId):

        # Assigning the UserId to a variable so it can be accessed by other functions inside this Class
        self.userId = userId

        # Assigning 'clientDetails' table name into a variable
        self.tableName = "clientDetails"

        # Connecting to the User's Database, if not exists creating a new database for each user
        self.inVoice_DB = sqlite3.connect(
            f"{UserCredential.DB_directory}/{userId}.db")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'clientDetails' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            clientName TEXT NOT NULL,
                            emailId CHAR NOT NULL,
                            contactNumber INTEGER NOT NULL,
                            addressLine1 TEXT NOT NULL,
                            addressLine2 TEXT,
                            addressLine3 TEXT,
                            cityName TEXT NOT NULL,
                            pinCode TEXT NOT NULL,
                            customerNote TEXT
                        )""")

    '''ADD CLIENT'''

    # Function for creating new record in the 'clientDetails' table
    def add(self, clientName, emailId, contactNumber, addressLine1, addressLine2, addressLine3, cityName, pinCode, customerNote):

        # Inserting a new record to the table
        self.co.execute(f"INSERT INTO {self.tableName} VALUES (:clientName, :emailId, :contactNumber, :addressLine1, :addressLine2, :addressLine3, :cityName, :pinCode, :customerNote)",
                        {
                            "clientName": clientName,
                            "emailId": emailId,
                            "contactNumber": contactNumber,
                            "addressLine1": addressLine1,
                            "addressLine2": addressLine2,
                            "addressLine3": addressLine3,
                            "cityName": cityName,
                            "pinCode": pinCode,
                            "customerNote": customerNote
                        })

    '''GET CLIENT'''

    # Function for getting specific record(client) from the 'clientDetails' table
    def get(self, clientId):

        # Querying the 'clientDetails' table, and picking the matched record(client)
        self.co.execute(
            f"SELECT * FROM {self.tableName} WHERE oid = '{clientId}'")

        # Assigning the queryed record(client) into a variable to return as response
        records = self.co.fetchone()

        return records

    '''UPDATE CLIENT'''

    # Function for updating specific record(client) in the 'clientDetails' table
    def update(self, clientId, clientName, emailId, contactNumber, addressLine1, addressLine2, addressLine3, cityName, pinCode, customerNote):

        # Updating a specific record(client) in the 'clientDetails' table using its oid
        self.co.execute(f"UPDATE {self.tableName} SET clientName = :clientName, emailId = :emailId, contactNumber = :contactNumber, addressLine1 = :addressLine1, addressLine2 = :addressLine2, addressLine3 = :addressLine3, cityName = :cityName, pinCode = :pinCode, customerNote = :customerNote WHERE oid = :uniqueId",
                        {
                            "clientName": clientName,
                            "emailId": emailId,
                            "contactNumber": contactNumber,
                            "addressLine1": addressLine1,
                            "addressLine2": addressLine2,
                            "addressLine3": addressLine3,
                            "cityName": cityName,
                            "pinCode": pinCode,
                            "customerNote": customerNote,

                            "uniqueId": clientId
                        })

    '''GET ALL CLIENTS'''

    # Function for getting all the records(clients) from the 'clientDetails' table
    def all(self):

        # Querying the 'clientDetails' table, and picking all the available records(clients)
        self.co.execute(f"SELECT *, oid FROM {self.tableName}")

        # Assigning the queryed records(clients) into a variable to return as response
        records = self.co.fetchall()

        return records

    '''DELETE CLIENT DETAILS'''

    # Function for deleting a specific record(client) from the 'clientDetails' table
    def delete(self, userId):

        # Deleting a specific record(client) from the 'clientDetails' table which matches the record(client)
        self.co.execute(
            f"DELETE FROM {self.tableName} WHERE oid = '{userId}'")

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Creating a class for doing all the CRUD operations in the 'productDetails' table on the inVoice User's Database
class ProductDetail():

    """
    Class for Creating a New Table in the User's Database,
    which contains all the Products created by the User for using it later on creating InVoice.
    """

    def __init__(self, userId):

        # Assigning the UserId to a variable so it can be accessed by other functions inside this Class
        self.userId = userId

        # Assigning 'productDetails' table name into a variable
        self.tableName = "productDetails"

        # Connecting to the User's Database, if not exists creating a new database for each user
        self.inVoice_DB = sqlite3.connect(
            f"{UserCredential.DB_directory}/{userId}.db")

        # Creating a cursor
        self.co = self.inVoice_DB.cursor()

        # Creating a 'productDetails' table if not exists
        self.co.execute(f"""CREATE TABLE IF NOT EXISTS {self.tableName} (
                            productName TEXT NOT NULL UNIQUE,
                            productMRP REAL,
                            quantity INTEGER NOT NULL,
                            purchaseRate REAL,
                            salesRate REAL NOT NULL,
                            reOrderQuantity INTEGER NOT NULL
                        )""")

    '''ADD PRODUCT'''

    # Function for creating new record (product) in the 'productDetails' table
    def add(self, productName, productMRP, quantity, purchaseRate, salesRate, reOrderQuantity):

        # Inserting a new record (product) to the table
        self.co.execute(f"INSERT INTO {self.tableName} VALUES (:productName, :productMRP, :quantity, :purchaseRate, :salesRate, :reOrderQuantity)",
                        {
                            "productName": productName,
                            "productMRP": productMRP,
                            "quantity": quantity,
                            "purchaseRate": purchaseRate,
                            "salesRate": salesRate,
                            "reOrderQuantity": reOrderQuantity
                        })

    '''UPDATE PRODUCT'''

    # Function for updating specific record(product) in the 'productDetails' table
    def update(self, productId, productName, productMRP, quantity, purchaseRate, salesRate, reOrderQuantity, getBy="oid"):

        # Updating a specific record(product) in the 'productDetails' table using its oid
        self.co.execute(f"UPDATE {self.tableName} SET productName = :productName, productMRP = :productMRP, quantity = :quantity, purchaseRate = :purchaseRate, salesRate = :salesRate, reOrderQuantity = :reOrderQuantity WHERE {getBy} = :uniqueId",
                        {
                            "productName": productName,
                            "productMRP": productMRP,
                            "quantity": quantity,
                            "purchaseRate": purchaseRate,
                            "salesRate": salesRate,
                            "reOrderQuantity": reOrderQuantity,

                            "uniqueId": productId
                        })

    '''GET A PRODUCT BY NAME'''

    # Function for getting a specified record(product) from the 'productDetails' table
    def get(self, filterby, productName):

        # Querying the 'productDetails' table, and picking the matched record(product)
        self.co.execute(
            f"SELECT * FROM {self.tableName} WHERE {filterby} = '{productName}'")

        # Assigning the queryed record(product) into a variable to return as response
        records = self.co.fetchone()

        return records

    '''GET ALL PRODUCTS'''

    # Function for getting all the records(products) from the 'productDetails' table
    def all(self):

        # Querying the 'productDetails' table, and picking all the available records(products)
        self.co.execute(f"SELECT *, oid FROM {self.tableName}")

        # Assigning the queryed records(products) into a variable to return as response
        records = self.co.fetchall()

        return records

    '''DELETE PRODUCT DETAILS'''

    # Function for deleting a specific record(product) from the 'productDetails' table
    def delete(self, userId):

        # Deleting a specific record(product) from the 'productDetails' table which matches the record(product)
        self.co.execute(
            f"DELETE FROM {self.tableName} WHERE oid = '{userId}'")

    '''CLOSE DB'''

    # Function for Closing the connection to Database
    def closeDB(self):

        # Committing the changes to the database
        self.inVoice_DB.commit()

        # Closing the connection with the database
        self.inVoice_DB.close()


# Function for creating the InVoice in '.docx' file format
def wordDocGenerator(inVoiceId, clientData, purchaseDate, dueDate, currencySymbol, productDetails, subTotal, calculatedTAX, totalAmount, balanceAmount, customerMessage):

    # Function for adding a empty line
    def emptyLines(count):

        for _ in range(count):

            linespace_style = document.styles["Body Text"]
            linespace_style.font.size = Pt(10)
            document.add_paragraph(style=linespace_style).add_run("")

    # >>

    # Opening a new Word document for storing the InVoice details as human readable data
    document = Document()

    # -------------------------------
    #           Headings
    # -------------------------------

    # Creating the Main heading for the 'InVoice' document and Aligning it to the center
    mainheading = document.add_heading()
    mainheading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Settings custom margins for the 'InVoice' document
    sections = document.sections
    section = sections[0]
    section.top_margin = Inches(0.04)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.8)

    # Adding the first content and styling of the Main heading
    run = mainheading.add_run("in")
    run.font.size = Pt(55)
    run.font.name = "Magneto"
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x40)

    # Adding the first content and styling of the Main heading
    run = mainheading.add_run("Voice")
    run.font.size = Pt(12)
    run.font.name = "Matura MT Script Capitals"
    run.font.color.rgb = RGBColor(0x46, 0x46, 0x46)

    # Adding an empty line
    emptyLines(1)

    # -------------------------------
    #           InVoice-Id
    # -------------------------------

    # Creating the template for inVoice number
    inVoiceId_container = document.add_paragraph("id:  ")
    inVoiceId_container.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Filling the inVoice number template with the inVoiceId's value passed through
    inVoiceNumber = inVoiceId_container.add_run(f"{inVoiceId}")
    inVoiceNumber.font.name = "Consolas"
    inVoiceNumber.font.size = Pt(13)
    inVoiceNumber.font.bold = True

    # -------------------------------
    #        Client Details
    # -------------------------------

    # Creating the template for Client Details
    clientdetails_container = document.add_paragraph("")

    # Filling the Client Details template with the client's name, address and phone
    client_name = clientdetails_container.add_run(
        f"{clientData[0]}\n{clientData[3]}\n{clientData[4]}\n{clientData[6]}\n{clientData[5]}\n\n{clientData[2]}")
    client_name.font.name = "Times New Roman"

    # -------------------------------
    #           Due Date
    # -------------------------------

    # Creating the template for Due Date
    duedate_container = document.add_paragraph("Due Date :  ")
    duedate_container.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Filling the Due Date template with the due date
    due_date = duedate_container.add_run(f"{dueDate}")
    due_date.font.size = Pt(13)
    due_date.font.name = "Times New Roman"

    # Adding an empty line
    emptyLines(1)

    # -------------------------------
    #         Product Details
    # -------------------------------

    # Creating a table for holding the product purchased details
    product_table = document.add_table(1, 5)

    # Creating the Table Header
    heading_cells = product_table.rows[0].cells

    # Populating the Table Header
    heading_cells[0].text = "No"
    heading_cells[1].text = "Description"
    heading_cells[2].text = "Quantity"
    heading_cells[3].text = f"Rate {currencySymbol}"
    heading_cells[4].text = f"Amount {currencySymbol}"

    # Aligning and Styling the names in Table Header
    heading_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading_cells[1].paragraphs[0].runs[0].font.bold = True
    heading_cells[2].paragraphs[0].runs[0].font.bold = True
    heading_cells[3].paragraphs[0].runs[0].font.bold = True
    heading_cells[4].paragraphs[0].runs[0].font.bold = True
    heading_cells[1].paragraphs[0].runs[0].font.size = Pt(13)
    heading_cells[2].paragraphs[0].runs[0].font.size = Pt(13)
    heading_cells[3].paragraphs[0].runs[0].font.size = Pt(13)
    heading_cells[4].paragraphs[0].runs[0].font.size = Pt(13)
    product_table.rows[0].height = Inches(0.6)

    # Populating the product details inside the table
    for detail in productDetails:

        # Creating a new row for each product
        cells = product_table.add_row().cells

        # Filling the content for each field of the row
        cells[0].text = str(detail[0])
        cells[1].text = detail[1]
        cells[2].text = str(detail[2])
        cells[3].text = str(detail[3])
        cells[4].text = str(detail[4])

        # Aligning and Styling the each row
        cells[1].width = Inches(2)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Adding an empty line
    emptyLines(1)

    # -------------------------------
    #           Sub Total
    # -------------------------------

    # Creating the template for Sub Total
    subtotal_container = document.add_paragraph("Sub Total :  ")
    subtotal_container.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Filling the Sub Total template with the value of sub total
    sub_total = subtotal_container.add_run(f"{subTotal}")
    sub_total.font.size = Pt(14)
    sub_total.font.name = "Times New Roman"
    sub_total.font.bold = True

    # -------------------------------
    #            Sales Tax
    # -------------------------------

    # Creating the template for Sales Tax
    salestax_container = document.add_paragraph("Sales Tax :  ")
    salestax_container.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Filling the Sales Tax template with the value of sales tax
    sales_tax = salestax_container.add_run(f"{calculatedTAX}")
    sales_tax.font.size = Pt(13)
    sales_tax.font.name = "Times New Roman"
    sales_tax.font.bold = True

    # Adding an empty line
    emptyLines(1)

    # -------------------------------
    #          Total Amount
    # -------------------------------

    # Creating the template for Total Amount
    totalamount_container = document.add_paragraph("Total Amount :  ")
    totalamount_container.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Filling the Total Amount template with the value of total amount
    total_amount = totalamount_container.add_run(
        f"{totalAmount} {currencySymbol}")
    total_amount.font.size = Pt(15)
    total_amount.font.name = "Times New Roman"
    total_amount.font.bold = True

    # Adding an empty line
    emptyLines(2)

    # -------------------------------
    #       Customer Message
    # -------------------------------

    # Creating the template for Customer Message
    customermsg_container = document.add_paragraph("")
    customermsg_container.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filling the Customer Message template with the value of Customer Message
    customer_msg = customermsg_container.add_run(f"~ {customerMessage} ~")
    customer_msg.font.size = Pt(13)
    customer_msg.font.name = "Times New Roman"
    customer_msg.font.bold = True

    # -------------------------------
    #           Balance Amount
    # -------------------------------

    # Adding an empty line
    emptyLines(1)

    # Creating the template for Balance Amount
    balanceamount_container = document.add_paragraph("Balance Amount :  ")
    balanceamount_container.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Setting the value of Balance Amount by checking the Balance Amount is zero
    to_pay = "Nil" if (int(balanceAmount) ==
                       0) else f"{balanceAmount} {currencySymbol}"

    # Filling the Balance Amount template with the value of balance amount
    balance_amount = balanceamount_container.add_run(f"{to_pay}")
    balance_amount.font.size = Pt(13)
    balance_amount.font.name = "Times New Roman"
    balance_amount.font.bold = True

    # Setting the value of payment status by checking the Balance Amount
    paymentStatus = "Paid" if (int(balanceAmount) == 0) else "Pending"

    # Settings the Date purchase and Payment status as content for footer
    footerDate = section.footer.paragraphs[0]
    footerDate.text = f"Dated On :  {purchaseDate}\t\tPayment Status :  {paymentStatus}"
    footerDate.style = document.styles["Footer"]

    # Generating the name of the document from the inVoice id
    pathToDOCXFile = f"{inVoiceId}.docx"

    count = 0
    while (os.path.exists(pathToDOCXFile)):

        count += 1
        # Generating the new name of the document from the inVoice id
        pathToDOCXFile = f"{inVoiceId}_{count}.docx"

    # Saving the document as in the generated name from inVoice id
    document.save(pathToDOCXFile)
